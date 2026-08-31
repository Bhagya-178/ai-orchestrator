"""
Phase 3: RAG Service

Handles document ingestion (PDF, DOCX, TXT), embedding generation via Ollama,
vector storage in Qdrant, and retrieval for chat augmentation.
"""

import asyncio
import logging
import re
import uuid
from pathlib import Path
from typing import Any

import pymupdf
from docx import Document as DocxDocument
from qdrant_client import AsyncQdrantClient
from qdrant_client.http import models as qdrant_models
from sqlalchemy import select, delete
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import (
    EMBEDDING_MODEL,
    OLLAMA_URL,
    QDRANT_COLLECTION,
    QDRANT_URL,
)
from app.database.models import Document, DocumentChunk
from app.ollama_client import ollama

logger = logging.getLogger(__name__)

# -- Configurable defaults -------------------------------------------------
# TOKEN-BASED CHUNKING (not character-based)
# Approximate: 1 token ≈ 4 characters for English text
DEFAULT_CHUNK_SIZE_TOKENS = 512  # ~2048 characters; good for 8K context windows
DEFAULT_CHUNK_OVERLAP_TOKENS = 64  # ~256 characters of overlap between chunks
EMBEDDING_BATCH_CONCURRENCY = 5

# Helper: Convert between tokens and characters
CHARS_PER_TOKEN = 4  # Rough approximation for English


def estimate_tokens(text: str) -> int:
    """Estimate token count using character approximation."""
    return len(text) // CHARS_PER_TOKEN


def estimate_chars_for_tokens(tokens: int) -> int:
    """Convert token count to approximate character count."""
    return tokens * CHARS_PER_TOKEN


class RAGService:
    """RAG pipeline: parse -> chunk -> embed -> store -> retrieve."""

    def __init__(self):
        self.client: AsyncQdrantClient | None = None
        self.collection_name = QDRANT_COLLECTION
        self._initialized = False
        self._available = False

    @property
    def is_available(self) -> bool:
        """Whether Qdrant is reachable and the collection is ready."""
        return self._available

    # ------------------------------------------------------------------
    # Qdrant initialisation  (lazy, async)
    # ------------------------------------------------------------------
    async def _init_client(self, required: bool = False):
        """Lazy-connect to Qdrant.

        Args:
            required: If True raise on failure (write ops like ingest/delete).
                      If False silently disable RAG (read ops like search).
        """
        if self._initialized:
            return
        try:
            self.client = AsyncQdrantClient(url=QDRANT_URL)
            await self._ensure_collection()
            self._initialized = True
            self._available = True
            logger.info("Connected to Qdrant at %s", QDRANT_URL)
        except Exception:
            self._available = False
            self._initialized = False
            self.client = None
            if required:
                logger.exception("Qdrant unreachable at %s", QDRANT_URL)
                raise RuntimeError(
                    f"Qdrant is not reachable at {QDRANT_URL}. "
                    "Please start Qdrant before uploading documents."
                )
            logger.warning(
                "Qdrant not reachable at %s — RAG features disabled.", QDRANT_URL
            )

    async def _ensure_collection(self):
        """Create collection if it doesn't exist."""
        try:
            collections = await self.client.get_collections()
            if not any(c.name == self.collection_name for c in collections.collections):
                await self.client.create_collection(
                    collection_name=self.collection_name,
                    vectors_config=qdrant_models.VectorParams(
                        size=1024,  # bge-m3 embedding dimension
                        distance=qdrant_models.Distance.COSINE,
                    ),
                )
                logger.info("Created Qdrant collection '%s'", self.collection_name)
        except Exception:
            logger.exception(
                "Failed to ensure Qdrant collection '%s'", self.collection_name
            )
            raise

    # ------------------------------------------------------------------
    # Document Parsing
    # ------------------------------------------------------------------
    async def parse_document(self, file_path: str, content_type: str) -> list[dict]:
        """Parse document and return list of {content, page_num, ...} dicts."""
        ext = Path(file_path).suffix.lower()
        try:
            if ext == ".pdf":
                return self._parse_pdf(file_path)
            elif ext in (".docx", ".doc"):
                return self._parse_docx(file_path)
            elif ext in (".txt", ".md"):
                return self._parse_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {ext}")
        except ValueError:
            raise
        except Exception:
            logger.exception("Failed to parse document: %s", file_path)
            raise ValueError(f"Error reading file '{Path(file_path).name}'")

    def _parse_pdf(self, file_path: str) -> list[dict]:
        """Extract text from PDF with page numbers."""
        chunks = []
        doc = pymupdf.open(file_path)
        try:
            for page_num, page in enumerate(doc, 1):
                text = page.get_text().strip()
                if text:
                    chunks.append({
                        "content": text,
                        "page_num": page_num,
                        "metadata": {"source": "pdf", "page": page_num},
                    })
        finally:
            doc.close()
        return chunks

    def _parse_docx(self, file_path: str) -> list[dict]:
        """Extract text from DOCX with section-level granularity.

        Paragraphs are grouped into sections of roughly DEFAULT_CHUNK_SIZE_TOKENS
        tokens (~2KB per section) so downstream chunk_text() works with reasonably-sized
        inputs and metadata stays meaningful.
        """
        chunks = []
        doc = DocxDocument(file_path)

        section_paragraphs: list[str] = []
        section_char_count = 0
        section_num = 1
        target_chars = estimate_chars_for_tokens(DEFAULT_CHUNK_SIZE_TOKENS)

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            section_paragraphs.append(text)
            section_char_count += len(text)

            if section_char_count >= target_chars:
                chunks.append({
                    "content": "\n".join(section_paragraphs),
                    "page_num": section_num,
                    "metadata": {"source": "docx", "section": section_num},
                })
                section_paragraphs = []
                section_char_count = 0
                section_num += 1

        if section_paragraphs:
            chunks.append({
                "content": "\n".join(section_paragraphs),
                "page_num": section_num,
                "metadata": {"source": "docx", "section": section_num},
            })

        return chunks

    def _parse_text(self, file_path: str) -> list[dict]:
        """Read plain text file."""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read().strip()
        if content:
            return [{
                "content": content,
                "page_num": 1,
                "metadata": {"source": "text"},
            }]
        return []

    # ------------------------------------------------------------------
    # Chunking  (sentence-aware)
    # ------------------------------------------------------------------
    @staticmethod
    def _split_sentences(text: str) -> list[str]:
        """Split text into sentences using a regex heuristic."""
        parts = re.split(r'(?<=[.!?])\s+', text)
        return [p for p in parts if p.strip()]

    def chunk_text(
        self,
        text: str,
        chunk_size_tokens: int = DEFAULT_CHUNK_SIZE_TOKENS,
        overlap_tokens: int = DEFAULT_CHUNK_OVERLAP_TOKENS,
    ) -> list[str]:
        """Split text into overlapping chunks, preferring sentence boundaries.

        Uses TOKEN-BASED sizing (not character-based) for better LLM performance.
        Greedily packs whole sentences up to chunk_size_tokens (~2KB per chunk).
        Falls back to character-level splitting only when a single
        sentence exceeds chunk_size_tokens.
        
        Args:
            text: Input text to chunk
            chunk_size_tokens: Target chunk size in tokens (~2048 chars for 512 tokens)
            overlap_tokens: Overlap between chunks in tokens (~256 chars for 64 tokens)
        """
        chunk_size_chars = estimate_chars_for_tokens(chunk_size_tokens)
        overlap_chars = estimate_chars_for_tokens(overlap_tokens)
        
        sentences = self._split_sentences(text)
        if not sentences:
            return [text] if text.strip() else []

        chunks: list[str] = []
        current_sentences: list[str] = []
        current_len = 0

        for sentence in sentences:
            sentence_len = len(sentence)

            # Oversized single sentence → hard-split by characters.
            if sentence_len > chunk_size_chars:
                if current_sentences:
                    chunks.append(" ".join(current_sentences))
                    current_sentences = []
                    current_len = 0

                start = 0
                while start < sentence_len:
                    end = min(start + chunk_size_chars, sentence_len)
                    chunks.append(sentence[start:end])
                    start += chunk_size_chars - overlap_chars
                continue

            # Would adding this sentence exceed the budget?
            if current_len + sentence_len + (1 if current_sentences else 0) > chunk_size_chars:
                chunks.append(" ".join(current_sentences))

                # Carry tail sentences as overlap context.
                overlap_sentences: list[str] = []
                overlap_len = 0
                for s in reversed(current_sentences):
                    if overlap_len + len(s) + (1 if overlap_sentences else 0) > overlap_chars:
                        break
                    overlap_sentences.insert(0, s)
                    overlap_len += len(s) + (1 if len(overlap_sentences) > 1 else 0)

                current_sentences = overlap_sentences
                current_len = (
                    sum(len(s) for s in current_sentences)
                    + max(0, len(current_sentences) - 1)
                )

            current_sentences.append(sentence)
            current_len += sentence_len + (1 if len(current_sentences) > 1 else 0)

        if current_sentences:
            chunks.append(" ".join(current_sentences))

        return chunks

    # ------------------------------------------------------------------
    # Embeddings  (concurrent batching)
    # ------------------------------------------------------------------
    async def get_embedding(self, text: str) -> list[float]:
        """Generate embedding using Ollama embedding model."""
        try:
            response = await ollama.embeddings(
                model=EMBEDDING_MODEL,
                prompt=text,
            )
            embedding = response.get("embedding", [])
            if not embedding:
                logger.warning("Empty embedding returned (text len=%d)", len(text))
            return embedding
        except Exception:
            logger.exception("Failed to generate embedding (text len=%d)", len(text))
            return []

    async def get_embeddings_batch(self, texts: list[str]) -> list[list[float]]:
        """Generate embeddings concurrently with bounded parallelism."""
        semaphore = asyncio.Semaphore(EMBEDDING_BATCH_CONCURRENCY)

        async def _embed(text: str) -> list[float]:
            async with semaphore:
                return await self.get_embedding(text)

        return await asyncio.gather(*[_embed(t) for t in texts])

    # ------------------------------------------------------------------
    # Ingestion Pipeline
    # ------------------------------------------------------------------
    async def ingest_document(
        self,
        db: AsyncSession,
        file_path: str,
        filename: str,
        content_type: str,
        file_size: int,
        session_id: str | None = None,
    ) -> Document:
        """Full ingestion: parse -> chunk -> embed -> store in Qdrant + Postgres."""

        # 1. Parse document
        parsed_chunks = await self.parse_document(file_path, content_type)

        if not parsed_chunks:
            raise ValueError(f"No content could be extracted from '{filename}'")

        # 2. Create Document record
        doc = Document(
            filename=filename,
            content_type=content_type,
            file_size=file_size,
            session_id=session_id,
            doc_metadata={"total_chunks": len(parsed_chunks)},
        )
        db.add(doc)
        await db.flush()  # get doc.id

        # 3. Process each parsed section
        all_chunks = []
        for section in parsed_chunks:
            text_chunks = self.chunk_text(section["content"])
            for idx, chunk_content in enumerate(text_chunks):
                all_chunks.append({
                    "content": chunk_content,
                    "chunk_index": len(all_chunks),
                    "page_num": section.get("page_num", 1),
                    "metadata": {**section.get("metadata", {}), "chunk_idx": idx},
                })

        # Update metadata with actual chunk count after splitting.
        doc.doc_metadata = {"total_chunks": len(all_chunks)}

        # 4. Generate embeddings (concurrent)
        embeddings = await self.get_embeddings_batch(
            [c["content"] for c in all_chunks]
        )

        # 5. Store in Qdrant + Postgres
        await self._init_client(required=True)
        points = []
        for chunk_data, embedding in zip(all_chunks, embeddings):
            if not embedding:
                logger.warning(
                    "Skipping chunk %d — empty embedding", chunk_data["chunk_index"]
                )
                continue

            point_id = str(uuid.uuid4())

            points.append(qdrant_models.PointStruct(
                id=point_id,
                vector=embedding,
                payload={
                    "document_id": str(doc.id),
                    "content": chunk_data["content"],
                    "chunk_index": chunk_data["chunk_index"],
                    "page_num": chunk_data["page_num"],
                    "metadata": chunk_data["metadata"],
                },
            ))

            db.add(DocumentChunk(
                document_id=doc.id,
                chunk_index=chunk_data["chunk_index"],
                content=chunk_data["content"],
                qdrant_point_id=point_id,
                chunk_metadata=chunk_data["metadata"],
            ))

        if not points and all_chunks:
            # If we had chunks but couldn't generate ANY embeddings, fail the upload.
            logger.error("Failed to generate embeddings for all %d chunks.", len(all_chunks))
            raise ValueError(f"Failed to generate embeddings for document. Is the embedding model '{EMBEDDING_MODEL}' pulled in Ollama?")

        # Batch upsert to Qdrant
        if points:
            try:
                await self.client.upsert(
                    collection_name=self.collection_name,
                    points=points,
                )
            except Exception:
                logger.exception("Qdrant upsert failed for document %s", doc.id)
                raise

        await db.commit()
        await db.refresh(doc)
        logger.info(
            "Ingested '%s': %d chunks, %d points stored",
            filename, len(all_chunks), len(points),
        )
        return doc

    # ------------------------------------------------------------------
    # Retrieval
    # ------------------------------------------------------------------
    async def search(
        self,
        query: str,
        limit: int = 5,
        score_threshold: float = 0.7,
        document_ids: list[str] | None = None,
    ) -> list[dict]:
        """Search for relevant chunks. Returns [] if Qdrant is unavailable."""
        await self._init_client(required=False)
        if not self._available:
            return []

        query_embedding = await self.get_embedding(query)
        if not query_embedding:
            logger.warning("Empty query embedding — returning no results")
            return []

        # Build filter if document_ids provided
        query_filter = None
        if document_ids:
            query_filter = qdrant_models.Filter(
                must=[
                    qdrant_models.FieldCondition(
                        key="document_id",
                        match=qdrant_models.MatchAny(any=document_ids),
                    )
                ]
            )

        try:
            results = await self.client.search(
                collection_name=self.collection_name,
                query_vector=query_embedding,
                limit=limit,
                score_threshold=score_threshold,
                query_filter=query_filter,
                with_payload=True,
            )
        except Exception:
            logger.exception("Qdrant search failed")
            return []

        return [
            {
                "content": hit.payload["content"],
                "score": hit.score,
                "document_id": hit.payload["document_id"],
                "chunk_index": hit.payload["chunk_index"],
                "page_num": hit.payload.get("page_num"),
                "metadata": hit.payload.get("metadata", {}),
            }
            for hit in results
        ]

    async def get_raw_chunks(self, db: AsyncSession, document_ids: list[str], limit: int = 5) -> list[dict]:
        """Directly fetch the first few chunks from Postgres as a fallback when search fails."""
        if not document_ids:
            return []
        
        query = (
            select(DocumentChunk)
            .where(DocumentChunk.document_id.in_(document_ids))
            .order_by(DocumentChunk.document_id, DocumentChunk.chunk_index)
            .limit(limit)
        )
        
        result = await db.execute(query)
        chunks = result.scalars().all()
        
        return [
            {
                "content": c.content,
                "score": 0.0,
                "document_id": str(c.document_id),
                "chunk_index": c.chunk_index,
                "page_num": c.chunk_metadata.get("page", 1),
                "metadata": c.chunk_metadata,
            }
            for c in chunks
        ]

    # ------------------------------------------------------------------
    # Document Management
    # ------------------------------------------------------------------
    async def list_documents(
        self, db: AsyncSession, session_id: str | None = None
    ) -> list[Document]:
        """List all documents, optionally filtered by session."""
        query = select(Document)
        if session_id:
            query = query.where(Document.session_id == session_id)
        result = await db.execute(query)
        return result.scalars().all()

    async def delete_document(self, db: AsyncSession, document_id: str) -> bool:
        """Delete document and its chunks from both stores."""
        await self._init_client(required=True)

        # Delete from Qdrant
        try:
            await self.client.delete(
                collection_name=self.collection_name,
                points_selector=qdrant_models.FilterSelector(
                    filter=qdrant_models.Filter(
                        must=[
                            qdrant_models.FieldCondition(
                                key="document_id",
                                match=qdrant_models.MatchValue(value=document_id),
                            )
                        ]
                    )
                ),
            )
        except Exception:
            logger.exception("Qdrant delete failed for document %s", document_id)
            # Continue to Postgres cleanup even if Qdrant fails.

        # Delete from Postgres (cascades to chunks)
        try:
            doc_uuid = uuid.UUID(document_id)
            # Delete chunks first to avoid foreign key constraint violation
            await db.execute(
                delete(DocumentChunk).where(DocumentChunk.document_id == doc_uuid)
            )
            result = await db.execute(
                select(Document).where(Document.id == doc_uuid)
            )
            doc = result.scalars().first()
            if doc:
                await db.delete(doc)
                await db.commit()
                return True
            # If doc not found but chunks deleted, we still commit
            await db.commit()
            return False
        except ValueError:
            # Invalid UUID
            return False
        except Exception as e:
            logger.error("Failed to delete document from Postgres: %s", e)
            await db.rollback()
            return False


rag_service = RAGService()